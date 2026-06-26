from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io


def get_patient_label(genre: str) -> str:
    if genre == 'feminin':
        return 'la patiente'
    elif genre == 'masculin':
        return 'le patient'
    return 'le-la patient·e'

def get_pronom(genre: str) -> str:
    if genre == 'feminin':
        return 'Elle'
    elif genre == 'masculin':
        return 'Il'
    return 'Il·elle'

def get_de_patient(genre: str) -> str:
    if genre == 'feminin':
        return 'de la patiente'
    elif genre == 'masculin':
        return 'du patient'
    return 'du-de la patient·e'

def get_symptomes_atm(symptomes) -> str:
    labels = {
        'douleursMachoire': 'douleurs de la mâchoire',
        'craquements': 'craquements',
        'deboitages': 'déboitages',
        'cephalees': 'céphalées',
        'troublesMastication': 'troubles de la mastication',
    }
    try:
        data = symptomes.model_dump()
    except Exception:
        data = symptomes.__dict__
    coches = [labels[k] for k, v in data.items() if v and k in labels]
    if not coches:
        return 'symptômes non précisés'
    return ', '.join(coches)

def val(texte: str) -> str:
    if not texte or not texte.strip():
        return 'Non évalué'
    return texte.strip()

def format_date_rapport(date_str_raw: str) -> str:
    mois = {
        1: 'janvier', 2: 'février', 3: 'mars', 4: 'avril',
        5: 'mai', 6: 'juin', 7: 'juillet', 8: 'août',
        9: 'septembre', 10: 'octobre', 11: 'novembre', 12: 'décembre'
    }
    try:
        date = datetime.strptime(date_str_raw, '%Y-%m-%d')
        return f"{date.day} {mois[date.month]} {date.year}"
    except Exception:
        now = datetime.now()
        return f"{now.day} {mois[now.month]} {now.year}"


def section_osteo_evaluee(form_data) -> bool:
    return any([
        form_data.instabiliteArticulaire,
        form_data.instabiliteATM,
        form_data.douleursInflammatoires,
        form_data.fracturesMultiples,
        form_data.syndromeDefileThoracique,
        form_data.piedsPats,
    ])

def section_generalisees_evaluee(form_data) -> bool:
    return any([
        form_data.douleursNociplastiques,
        form_data.douleursMusculaires,
    ])

def section_neuropathiques_evaluee(form_data) -> bool:
    return bool(form_data.douleursNeuropathiques)

def section_recommandations_evaluee(form_data) -> bool:
    try:
        recs = form_data.recommandations.model_dump()
    except Exception:
        recs = form_data.recommandations.__dict__
    return any(recs.values()) or bool(
        form_data.recommandationsTexteLibre and
        form_data.recommandationsTexteLibre.strip()
    )


def add_title(doc, text, level=1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13) if level == 1 else Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    para.space_before = Pt(14)
    para.space_after = Pt(6)
    return para

def add_paragraph(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    para.space_after = Pt(6)
    return para

def add_bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(11)
    para.space_after = Pt(4)
    return para

def add_notes_libres(doc, notes: str):
    if notes and notes.strip():
        para = doc.add_paragraph()
        run_label = para.add_run("Notes libres : ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_notes = para.add_run(notes.strip())
        run_notes.font.size = Pt(11)
        run_notes.italic = True
        para.space_after = Pt(6)

def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_report(patient_id: str, form_data, destinataire: str = '[Destinataire]') -> bytes:
    doc = Document()
    genre = form_data.genre or ''
    patient = get_patient_label(genre)
    pronom = get_pronom(genre)
    de_patient = get_de_patient(genre)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    date_str = format_date_rapport(form_data.dateConsultation or '')

    table_header = doc.add_table(rows=1, cols=2)
    for row in table_header.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    cell_left = table_header.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    run = p_left.add_run('[Adresse expéditeur]')
    run.italic = True
    run.font.size = Pt(11)

    cell_right = table_header.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_right.add_run(f'Lausanne, le {date_str}')
    run.font.size = Pt(11)

    doc.add_paragraph()

    para_dest = doc.add_paragraph()
    para_dest.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lignes = destinataire.split('\n')
    for i, ligne in enumerate(lignes):
        run = para_dest.add_run(ligne)
        run.font.size = Pt(11)
        run.italic = True
        if i < len(lignes) - 1:
            run.add_break()

    doc.add_paragraph()
    doc.add_paragraph()

    para_concerne = doc.add_paragraph()
    run = para_concerne.add_run(
        'Concerne : MME/MR [Nom Prénom], [dd/mm/yyyy], [adresse à compléter]'
    )
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '[Madame / Monsieur],')
    doc.add_paragraph()

    add_paragraph(doc,
        "Cette consultation avait pour objectif d'évaluer différents aspects cliniques fréquemment "
        "associés au syndrome d'Ehlers-Danlos hypermobile (SEDh) ou aux troubles du spectre de "
        "l'hypermobilité (HSD), afin d'orienter la prise en charge et les éventuelles investigations "
        "complémentaires."
    )
    add_paragraph(doc,
        "Vous trouverez ci-dessous une synthèse des éléments présentant un intérêt clinique ainsi "
        "que les recommandations qui en découlent."
    )
    doc.add_paragraph()

    add_title(doc, 'Douleurs ostéo-articulaires')

    if not section_osteo_evaluee(form_data):
        add_paragraph(doc,
            "La composante ostéo-articulaire des douleurs n'a pas été évaluée lors de cette consultation."
        )
    else:
        r = form_data.instabiliteArticulaire
        if r == 'oui':
            add_paragraph(doc,
                f"En présence d'une instabilité articulaire, l'utilisation de dispositifs compressifs "
                f"(vêtements compressifs, orthèses adaptées) ainsi qu'une physiothérapie ciblée sur la "
                f"proprioception, le contrôle moteur et la stabilisation articulaire doivent être proposées "
                f"afin d'améliorer la fonction et de réduire la fréquence des épisodes de luxation ou de "
                f"subluxation. Une liste de physiothérapeutes formés à la prise en charge du SEDh peut être "
                f"fournie si nécessaire."
            )
        elif r == 'non':
            add_paragraph(doc,
                f"{patient.capitalize()} ne présente pas d'instabilité articulaire avec luxations ou subluxations."
            )
        else:
            add_paragraph(doc,
                "La présence d'une instabilité articulaire avec luxations ou subluxations n'a pas été "
                "évaluée lors de cette consultation."
            )

        r = form_data.instabiliteATM
        symptomes_str = get_symptomes_atm(form_data.instabiliteATMSymptomes)
        if r == 'oui':
            pec = form_data.instabiliteATMPriseEnCharge
            physio = form_data.instabiliteATMPhysio
            chir = form_data.instabiliteATMChirurgie
            if pec == 'oui':
                if physio == 'oui' and chir == 'oui':
                    add_paragraph(doc,
                        f"{patient.capitalize()} présente une instabilité des ATM associée à {symptomes_str}. "
                        f"{pronom} bénéficie actuellement d'une physiothérapie spécifique ainsi que d'un "
                        f"suivi en chirurgie maxillo-faciale."
                    )
                elif physio == 'oui' and chir != 'oui':
                    add_paragraph(doc,
                        f"{patient.capitalize()} présente une instabilité des ATM associée à {symptomes_str}. "
                        f"{pronom} bénéficie actuellement d'une physiothérapie spécifique."
                    )
                elif physio != 'oui' and chir == 'oui':
                    add_paragraph(doc,
                        f"{patient.capitalize()} présente une instabilité des ATM associée à {symptomes_str}. "
                        f"{pronom} bénéficie actuellement d'un suivi en chirurgie maxillo-faciale sans "
                        f"physiothérapie spécifique à ce stade."
                    )
                else:
                    add_paragraph(doc,
                        f"{patient.capitalize()} présente une instabilité des ATM associée à {symptomes_str}. "
                        f"Une prise en charge est rapportée, sans précision concernant son type."
                    )
            elif pec == 'non':
                add_paragraph(doc,
                    f"{patient.capitalize()} présente {symptomes_str}, évocateurs d'une instabilité des "
                    f"articulations temporo-mandibulaires (ATM). Aucune prise en charge spécifique n'est "
                    f"actuellement rapportée. Une évaluation en chirurgie maxillo-faciale pourrait être "
                    f"proposée afin d'estimer le potentiel de réhabilitation, notamment par une "
                    f"physiothérapie spécifique."
                )
            else:
                add_paragraph(doc,
                    f"{patient.capitalize()} présente {symptomes_str}, évocateurs d'une instabilité des "
                    f"articulations temporo-mandibulaires (ATM). Les modalités de prise en charge n'ont "
                    f"pas été évaluées lors de cette consultation."
                )
        elif r == 'non':
            add_paragraph(doc,
                "Aucun signe spécifique en faveur d'une instabilité des articulations temporo-mandibulaires "
                "(ATM) n'a été identifié."
            )
        else:
            add_paragraph(doc,
                "La présence d'une instabilité des articulations temporo-mandibulaires (ATM) n'a pas été "
                "évaluée lors de cette consultation."
            )

        r = form_data.douleursInflammatoires
        type_inflam = val(form_data.typeInflammation)
        if r == 'oui':
            add_paragraph(doc,
                f"Les douleurs {de_patient} présentent une composante inflammatoire de type {type_inflam}. "
                f"Une évaluation par un rhumatologue serait indiquée afin de clarifier le diagnostic. "
                f"Si localisée, une approche interventionnelle par un antalgiste pourrait être considérée."
            )
        elif r == 'non':
            add_paragraph(doc, "Les douleurs ne semblent pas être de type inflammatoire ou dégénérative.")
        else:
            add_paragraph(doc,
                "La présence de douleurs avec une composante inflammatoire n'a pas été évaluée lors de "
                "cette consultation."
            )

        r = form_data.fracturesMultiples
        if r == 'oui':
            osteo = form_data.osteoporoseConnue
            traitement = form_data.traitementOsteoporose
            ouvert = form_data.patientOuvertTraitement
            if osteo == 'oui':
                if traitement == 'oui':
                    add_paragraph(doc, "Une ostéoporose est connue et prise en charge.")
                elif traitement == 'non':
                    if ouvert == 'oui':
                        add_paragraph(doc,
                            f"Compte tenu des antécédents de fractures multiples et/ou de fractures de "
                            f"fragilité, une évaluation de la densité minérale osseuse et des facteurs de "
                            f"risque de fragilité osseuse peut être discutée si elle n'a pas déjà été "
                            f"effectuée. Une prise en charge adaptée est recommandée en cas d'ostéoporose "
                            f"ou d'ostéopénie significative."
                        )
                    elif ouvert == 'non':
                        add_paragraph(doc,
                            f"{patient.capitalize()} ne souhaite pas mettre en place un traitement."
                        )
                    else:
                        add_paragraph(doc,
                            f"L'ouverture {de_patient} à l'idée d'un traitement n'a pas été évaluée lors "
                            f"de cette consultation."
                        )
                else:
                    add_paragraph(doc,
                        "Le traitement spécifique de l'ostéoporose en cours n'a pas été évalué lors de "
                        "cette consultation."
                    )
            elif osteo == 'non':
                add_paragraph(doc,
                    f"{patient.capitalize()} ne présente pas d'antécédents de fractures multiples et/ou "
                    f"de fractures survenues après un traumatisme mineur (fractures de fragilité)."
                )
            else:
                add_paragraph(doc,
                    "L'ostéoporose connue n'a pas été évaluée lors de cette consultation."
                )
        elif r == 'non':
            add_paragraph(doc,
                f"{patient.capitalize()} ne présente pas de signes de fragilité osseuse."
            )
        else:
            add_paragraph(doc,
                "Les antécédents de fractures multiples et/ou de fractures de fragilité n'ont pas été "
                "évalués lors de cette consultation."
            )

        r = form_data.syndromeDefileThoracique
        if r == 'oui':
            connu = form_data.syndromeDefileConnu
            pec = form_data.syndromeDefilePriseEnCharge
            if connu == 'oui':
                if pec == 'oui':
                    add_paragraph(doc, "Un SDT est connu et pris en charge.")
                elif pec == 'non':
                    add_paragraph(doc,
                        "Un SDT est connu mais aucune prise en charge n'est actuellement rapportée."
                    )
                else:
                    add_paragraph(doc,
                        "Un SDT est connu. La prise en charge n'a pas été évaluée lors de cette consultation."
                    )
            elif connu == 'non':
                add_paragraph(doc,
                    f"Les symptômes présentés par {patient} évoquent la présence d'un éventuel SDT. "
                    f"Une évaluation plus poussée par un angiologue serait indiquée."
                )
            else:
                add_paragraph(doc,
                    "La présence d'un syndrome du défilé thoracique n'a pas été évaluée lors de cette "
                    "consultation."
                )
        elif r == 'non':
            add_paragraph(doc,
                f"{patient.capitalize()} ne présente pas de signe évocateur d'un syndrome du défilé "
                f"thoracique (SDT)."
            )
        else:
            add_paragraph(doc,
                "La suspicion d'un syndrome du défilé thoracique (SDT) n'a pas été évaluée lors de "
                "cette consultation."
            )

        r = form_data.piedsPats
        podo = form_data.piedsPatsPodiologie
        if r == 'oui':
            if podo == 'oui':
                add_paragraph(doc,
                    f"{patient.capitalize()} a les pieds plats et bénéficie d'une prise en charge en "
                    f"podologie."
                )
            elif podo == 'non':
                add_paragraph(doc,
                    f"{patient.capitalize()} a les pieds plats et n'a pas encore de prise en charge. "
                    f"Une consultation chez un podologue spécialisé dans les semelles orthopédiques est "
                    f"recommandée."
                )
            else:
                add_paragraph(doc,
                    f"{patient.capitalize()} a les pieds plats. La prise en charge par un podologue n'a "
                    f"pas été évaluée lors de cette consultation."
                )
        elif r == 'non':
            add_paragraph(doc, f"{patient.capitalize()} n'a pas les pieds plats.")
        else:
            add_paragraph(doc,
                "La présence de pieds plats n'a pas été évaluée lors de cette consultation."
            )

        add_notes_libres(doc, form_data.notesOsteo)

    doc.add_paragraph()

    add_title(doc, 'Douleurs généralisées')

    if not section_generalisees_evaluee(form_data):
        add_paragraph(doc,
            "La composante des douleurs généralisées n'a pas été évaluée lors de cette consultation."
        )
    else:
        r = form_data.douleursNociplastiques
        desc = val(form_data.descriptionDouleursNociplastiques)
        if r == 'oui':
            add_paragraph(doc,
                f"Les douleurs {de_patient} ont une composante nociplastique (de type {desc}) qui demande "
                f"une prise en charge globale incluant une psycho-éducation de la douleur, une modulation "
                f"de l'hypervigilance et un ré-entraînement de la réponse de relaxation. Des approches "
                f"médicamenteuses spécifiques avec des anti-dépresseurs noradrénergiques ou des médicaments "
                f"gabaergiques (P.O.) ou possiblement des perfusions de lidocaïne ou kétamine dans un "
                f"centre d'antalgie peuvent être considérés. L'application de TENS avec une éducation "
                f"thérapeutique adaptée et associée à une stimulation du nerf vague peut être aidante."
            )
        elif r == 'non':
            add_paragraph(doc, "Les douleurs n'ont pas de composante nociplastique.")
        else:
            add_paragraph(doc,
                "La présence de douleurs primaires, nociplastiques ou de signes de sensibilisation "
                "centrale n'a pas été évaluée lors de cette consultation."
            )

        r = form_data.douleursMusculaires
        if r == 'oui':
            add_paragraph(doc,
                f"En présence de douleurs musculaires et d'une fatigabilité importante, il est recommandé "
                f"de rechercher des déficits ou troubles pouvant y contribuer (p.ex. dosages de vitamine D, "
                f"ferritine, TSH, troubles du sommeil). Une physiothérapie spécialisée et un "
                f"reconditionnement progressif sont nécessaires pour diminuer ce type de douleurs."
            )
        elif r == 'non':
            add_paragraph(doc,
                f"{patient.capitalize()} ne présente pas de douleur musculaire ou de fatigabilité excessive."
            )
        else:
            add_paragraph(doc,
                "La présence de douleurs musculaires ou d'une fatigabilité excessive n'a pas été évaluée "
                "lors de cette consultation."
            )

        add_notes_libres(doc, form_data.notesGeneralisees)

    doc.add_paragraph()


    add_title(doc, 'Douleurs neuropathiques')

    if not section_neuropathiques_evaluee(form_data):
        add_paragraph(doc,
            "La composante des douleurs neuropathiques n'a pas été évaluée lors de cette consultation."
        )
    else:
        r = form_data.douleursNeuropathiques
        if r == 'oui':
            diagnostic = form_data.diagnosticNeuropathique
            diag_texte = val(form_data.diagnosticNeuropathiqueTexte)
            pec = form_data.priseEnChargeNeuropathique
            if diagnostic == 'oui':
                if pec == 'oui':
                    add_paragraph(doc,
                        f"{patient.capitalize()} présente des douleurs à caractère neuropathique. "
                        f"Un diagnostic de {diag_texte} a été posé et bénéficie actuellement d'une "
                        f"prise en charge spécifique."
                    )
                elif pec == 'non':
                    add_paragraph(doc,
                        f"{patient.capitalize()} présente des douleurs à caractère neuropathique. "
                        f"Un diagnostic de {diag_texte} a été posé. Aucune prise en charge spécifique "
                        f"n'est actuellement rapportée. Un traitement spécifique de la douleur "
                        f"neuropathique pourrait être reconsidéré, éventuellement avec l'aide d'un "
                        f"antalgiste."
                    )
                else:
                    add_paragraph(doc,
                        f"{patient.capitalize()} présente des douleurs à caractère neuropathique. "
                        f"Un diagnostic de {diag_texte} a été posé. Les modalités de prise en charge "
                        f"n'ont pas été évaluées lors de cette consultation."
                    )
            elif diagnostic == 'non':
                add_paragraph(doc,
                    f"Afin d'évaluer la cause de ces douleurs à caractère neuropathique, une évaluation "
                    f"chez un neurologue est recommandée, qui jugera la pertinence d'exclure une atteinte "
                    f"des grosses fibres nerveuses (mono ou poly-neuropathie). Si cette condition n'est "
                    f"pas retenue, une évaluation d'une éventuelle neuropathie des petites fibres "
                    f"nerveuses pourrait être indiquée. Dans ce cas, le centre d'antalgie du CHUV propose "
                    f"une évaluation combinant un test de la fonction des fibres (QST) et une biopsie "
                    f"cutanée afin de déterminer la densité intraépidermale des petites fibres."
                )
                doc.add_paragraph()
                add_paragraph(doc,
                    f"En parallèle de l'élaboration du diagnostic, un traitement symptomatique doit être "
                    f"envisagé. Les approches thérapeutiques suivantes peuvent être considérées : TENS "
                    f"avec une éducation thérapeutique adaptée, traitements avec des agents anti-douleurs "
                    f"neuropathiques (antidépresseurs noradrénergiques, gabapentinoïdes), traitements "
                    f"topiques de lidocaïne ou à base d'amitriptyline et kétamine, perfusion intraveineuse "
                    f"de kétamine ou de lidocaïne."
                )
            else:
                add_paragraph(doc,
                    f"{patient.capitalize()} présente des douleurs à caractère neuropathique. L'existence "
                    f"d'un diagnostic spécifique n'a pas été évaluée lors de cette consultation."
                )
        elif r == 'non':
            add_paragraph(doc,
                f"{patient.capitalize()} ne présente pas de signes évocateurs de douleurs neuropathiques."
            )
        else:
            add_paragraph(doc,
                "La présence de douleurs neuropathiques n'a pas été évaluée lors de cette consultation."
            )

        add_notes_libres(doc, form_data.notesNeuropathiques)

    doc.add_paragraph()

    add_title(doc, 'Recommandations de traitements thérapeutiques')

    if not section_recommandations_evaluee(form_data):
        add_paragraph(doc,
            "Aucune recommandation thérapeutique n'a été sélectionnée lors de cette consultation.",
            italic=True
        )
    else:
        add_paragraph(doc, "Nous avons discuté des recommandations thérapeutiques suivantes :")

        recs = form_data.recommandations
        textes_recs = {
            'psychoeducation': (
                "une psychoéducation sur le syndrome d'Ehlers-Danlos (hEDS/HSD), la douleur et le "
                "reconditionnement comme pilier du traitement."
            ),
            'reconditionnement': (
                "un reconditionnement musculaire progressif pour la stabilisation articulaire, la "
                "diminution des contractures musculaires, la gestion de l'effort, et des adaptations "
                "des mouvements pour réduire les risques de luxations et blessures. Ceci est entrepris "
                "par des physiothérapeutes spécialisés. De plus, en auto-soins, les patients sont "
                "encouragés à trouver une activité physique progressive. Celles particulièrement "
                "recommandées sont : Pilates, qi gong, tai-chi, natation, marche, vélo couché. "
                "NB : Les ajustements chiropratiques et le yoga ne sont pas contre-indiqués, mais "
                "comme tous les autres traitements physiques, ils doivent être pratiqués de manière "
                "à éviter les subluxations ou luxations iatrogènes, et idéalement par des praticiens "
                "qui ont des bonnes connaissances de SEDh/TSH."
            ),
            'ergotherapie': (
                "l'ergothérapie qui peut adapter des aides techniques et accompagner les modifications "
                "de l'utilisation corporelle, l'ergonomie."
            ),
            'tens': (
                "le test accompagné d'un appareil de neurostimulation périphérique TENS +- stimulation "
                "du nerf vague."
            ),
            'antalgiques': (
                "les antalgiques sont à considérer et tester, en fonction des "
                "contre-indications/intolérances, y.c. : AINS, paracétamol, métamizole, "
                "gabapentinoïdes, antidépresseurs tricycliques, CBD/THC."
            ),
            'topiques': (
                "des traitements topiques peuvent être aidants tels que la capsaïcine, les AINS, "
                "la lidocaïne, l'arnica."
            ),
            'gestion': (
                "les techniques de cohérence cardiaque et techniques psycho-corporelles pour la "
                "gestion de la douleur, de l'anxiété et du sommeil."
            ),
            'chaleur': (
                "la chaleur mérite d'être testée (bains chauds, natation/exercices en piscine "
                "chauffée) pour favoriser la relaxation musculaire."
            ),
            'acupuncture': (
                "l'acupuncture peut avoir des bénéfices notamment pour des céphalées, des douleurs "
                "inflammatoires ou des symptômes de dysautonomie."
            ),
            'complementsAlimentaires': (
                "des compléments alimentaires peuvent être testés en l'absence d'intolérances tels "
                "que : les acides gras oméga-3 (huile de poisson), le curcuma, l'acide alpha-lipoïque, "
                "la carnitine (1 g 3 fois par jour)."
            ),
        }

        try:
            recs_dict = recs.model_dump()
        except Exception:
            recs_dict = recs.__dict__

        for key, texte in textes_recs.items():
            if recs_dict.get(key):
                add_bullet(doc, texte)

        if form_data.recommandationsTexteLibre and form_data.recommandationsTexteLibre.strip():
            doc.add_paragraph()
            add_paragraph(doc,
                "De plus, nous avons discuté de propositions thérapeutiques spécifiques :"
            )
            add_paragraph(doc, form_data.recommandationsTexteLibre)

    doc.add_paragraph()

    add_separator(doc)
    doc.add_paragraph()
    add_paragraph(doc,
        "Si vous avez des questions concernant ce rapport, nous restons à votre disposition. "
        "N'hésitez pas à nous contacter via les coordonnées fournies ci-dessus."
    )
    doc.add_paragraph()
    add_paragraph(doc, "Avec nos salutations les plus respectueuses.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()